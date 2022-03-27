from docx import Document
import requests
from bs4 import BeautifulSoup
import os
import comtypes.client
from PIL import Image
import glob
import tkinter as tk

from docx.shared import Cm


def main():
    gui = tk.Tk()

    webLbl = tk.Label(gui, text="Website Url:")
    locationLbl = tk.Label(gui, text="Save Location:")
    infoLabel = tk.Label(gui, text="")
    fileNameLbl = tk.Label(gui, text="FileName:")
    webLbl.grid(row=0, column=0, sticky=tk.W, pady=2)
    locationLbl.grid(row=1, column=0, sticky=tk.W, pady=2)
    fileNameLbl.grid(row=2, column=0, sticky=tk.W, pady=2)
    infoLabel.grid(row=3, column=0, sticky=tk.W, pady=2)

    webEntry = tk.Entry(gui)
    webEntry.config(width=50)
    locationEntry = tk.Entry(gui)
    locationEntry.config(width=50)
    fileNameEntry = tk.Entry(gui)
    fileNameEntry.config(width=15)
    webEntry.grid(row=0, column=1, pady=2)
    locationEntry.grid(row=1, column=1, pady=2)
    fileNameEntry.grid(row=2, column=1, sticky=tk.W, pady=2)

    convert = tk.IntVar(value=1)
    checkBox = tk.Checkbutton(gui, text='Convert to PDF', variable=convert, onvalue=1, offvalue=0, )
    checkBox.grid(row=4, column=0, sticky=tk.W, pady=2)

    clear = tk.IntVar(value=1)
    clearBox = tk.Checkbutton(gui, text='Clear Directory Previous Images', variable=clear, onvalue=1, offvalue=0)
    clearBox.grid(row=4, column=1, sticky=tk.W, pady=2)

    generateBtn = tk.Button(gui, text="Generate",
                            command=lambda: imageExtractor(webEntry.get(), locationEntry.get(), fileNameEntry.get(),
                                                           infoLabel,
                                                           generateBtn,
                                                           convert, clear))
    generateBtn.grid(row=5, column=1, pady=2)

    gui.mainloop()


def imageExtractor(url, folderPATH, fileName, infoLabel, generateBtn, convert, clear):
    changeButtonState(generateBtn, "disable")
    if not os.path.isdir(folderPATH):
        editLabelText(infoLabel, "O diretorio nao existe")
    elif not websiteExists(url):
        editLabelText(infoLabel, "Esse capitulo nao existe")
    else:
        os.chdir(folderPATH)
        if clear:
            editLabelText(infoLabel, "Clean das imagens desnecessarias")
            removeDirFiles('jpg', folderPATH, fileName)
            # editLabelText(label, "Download das imagens necessarias")
        downloadImages(url)
        if convert:
            # editLabelText(label, "A criar word")
            addImagesToWord(folderPATH, fileName)
            # editLabelText(label, "A converter word para pdf")
            wordToPDF(folderPATH, fileName)
            # editLabelText(label, "Feito!")
        if clear:
            removeDirFiles('jpg', folderPATH, fileName)
        editLabelText(infoLabel, "")
    changeButtonState(generateBtn, "normal")


def changeButtonState(button, state):
    button["state"] = state


def websiteExists(url):
    response = requests.get(url)
    if response.status_code == 200:
        return True
    else:
        return False


def editLabelText(label, newText):
    label.config(text=newText, highlightcolor="red")


def downloadImages(url):
    r = requests.get(url)
    soup = BeautifulSoup(r.text, 'html.parser')

    images = soup.find_all('img')
    i = 1100
    for image in images:
        link = image['src']
        with open("tempBerserk" + str(i) + '.jpg', 'wb') as f:
            im = requests.get(link)
            f.write(im.content)
        i = i + 1


def addImagesToWord(folderPATH, fileName):
    document = Document()
    updateMargins(document, folderPATH, fileName)
    document = Document(os.path.join(folderPATH, fileName + '.docx'))
    for path in glob.glob1(folderPATH, 'tempBerserk*.jpg'):
        img = Image.open(os.path.join(folderPATH, path))
        if img.size[0] > 1500:
            img = img.transpose(Image.ROTATE_270)
            img = img.resize((int(img.size[0] * 0.35), int(img.size[1] * 0.35)), Image.ANTIALIAS)
        else:
            img = img.resize((int(img.size[0] * 0.5), int(img.size[1] * 0.5)), Image.ANTIALIAS)
        img.save(os.path.join(folderPATH, path))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(os.path.join(folderPATH, path))
    document.save(os.path.join(folderPATH, fileName + '.docx'))


def updateMargins(document, folderPATH, fileName):
    sections = document.sections
    for section in sections:
        section.top_margin = Cm(0)
        section.left_margin = Cm(0.8)

    document.save(os.path.join(folderPATH, fileName + '.docx'))


def wordToPDF(folderPATH, fileName):
    in_file = os.path.abspath(os.path.join(folderPATH, fileName + '.docx'))
    out_file = os.path.abspath(os.path.join(folderPATH, fileName + '.pdf'))

    wdFormatPDF = 17

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(in_file)
    doc.SaveAs(out_file, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()


def removeDirFiles(fileType, folderPATH, fileName):
    for path in glob.glob1(folderPATH, 'tempBerserk*.' + fileType):
        os.remove(os.path.join(folderPATH, path))
    if os.path.isfile(os.path.join(folderPATH, fileName + '.docx')):
        os.remove(os.path.join(folderPATH, fileName + '.docx'))


if __name__ == '__main__':
    main()
